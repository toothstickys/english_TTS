from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    content = []
    #遍历读取doc文件中的文本
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i].text
        para = para.replace('\n', '')
        if '\xa0' in para:
            para = para.replace('\xa0', ' ')
        content.append(para)
    return content

def read_txt(file_path):
    content = []
    file_txt = open(file_path, "r", encoding='utf-8').readlines()
    for i in range(len(file_txt)):
        para = file_txt[i]
        para = para.replace('\n', '')
    #读取doc文件
    # doc = Document(argparsers.doc_path)
    # #遍历读取doc文件中的文本
    # for i in range(len(doc.paragraphs)):
    #     para = doc.paragraphs[i].text
        if '\xa0' in para:
            para = para.replace('\xa0', ' ')
        content.append(para)
    return content
